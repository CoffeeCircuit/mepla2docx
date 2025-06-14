"""
SJ Mepla to .docx file converter
Module reads .dat and .rep files from SJ Mepla
"""

__version__ = "0.1.0"

__all__ = ["document", "PATH", "FILES"]

import re
import argparse
from os import listdir
from os.path import isfile, isdir, join
from io import StringIO
from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_TAB_ALIGNMENT


def isnumber(val: str):
    """
    Returns true if a string is a valid number.
    """
    re_pattern = re.compile(
        r"(?:^[+|-]?[0-9]+$)|(?:^[+|-]?[0-9]+\.[0-9]+$)|(?:^[+|-]?[0-9]+\.[0-9]+[e|E]?[+|-]?[0-9]+$)"
    )
    if re_pattern.match(val):
        return True
    else:
        return False


def parse(line: str) -> tuple[list]:
    """
    Returns (left aligned) tabstops and words
    tabulated by spaces and/or underscores
    """
    padding: int = 0
    start: int = 0
    word = StringIO()
    words: list[str] = []
    tabstops: list[int] = []
    last_index: int = len(line) - 1
    for i, char in enumerate(line):
        if char in " _":
            padding += 1
        else:
            word.write(char)
            padding = 0
            if i != last_index:
                continue
        if word.getvalue():
            if padding == 1 and isnumber(word.getvalue()):
                words.append(word.getvalue().rstrip().strip("_"))
                word.truncate(0)
                word.seek(0)
                padding = 0
                continue
            if padding < 2:
                if i != last_index:
                    word.write(char)
            if padding == 2:
                words.append(word.getvalue().rstrip().strip("_"))
                word.truncate(0)
                word.seek(0)
        if i == last_index:
            if word.getvalue():
                words.append(word.getvalue().rstrip().strip("_"))
    for _word in words:
        p = line.find(_word, start)
        start = len(_word) + p + 1
        tabstops.append(p + 1)
    return tabstops, words


def parse_dat(dat_file: str, output_directory: str):
    """
    Parse sj_meplat.dat
    """
    document = Document()
    records = []
    with open(dat_file, "r", encoding="utf-8-sig") as fp:
        _FILE = tuple(fp.read().split("\n"))
    for line in _FILE:
        if line.startswith("00"):
            if records:
                for _record in records:
                    tabstops, words = parse(_record)
                    p = document.add_paragraph("\t" + "\t".join(words))
                    p.style = "No Spacing"
                    for tab in tabstops:
                        _t = p.paragraph_format.tab_stops
                        _t.add_tab_stop(
                            (Mm(tab * 1.75)), alignment=WD_TAB_ALIGNMENT.LEFT
                        )
                records = []
        if line.startswith("01"):
            if re.findall("={3,}", line):
                continue
            else:
                heading1 = line[2:].rstrip(":")
                document.add_heading(heading1, 1)
        if line.startswith("02"):
            heading2 = line[2:].rstrip(":")
            document.add_heading(heading2, 2)
        if line.startswith("03"):
            row = line.lstrip("03")
            records.append(row)
        if line.startswith("04"):
            row = line.lstrip("04")
            records.append(row)
    document.save(join(output_directory, "sj_mepla.dat.docx"))


def parse_rep(rep_file: str, output_directory: str):
    """
    Parse sj_meplat.rep
    """
    document = Document()
    with open(rep_file, "r", encoding="utf-8-sig") as fp:
        _FILE = tuple(fp.read().split("\n"))

    hrule = 0
    for line in _FILE:
        if line == "":
            p = document.add_paragraph()
            p.style.paragraph_format.space_before = 0
            p.style.paragraph_format.space_after = 0
            continue
        elif line.startswith("─"):
            hrule += 1
            if hrule == 1:
                heading = True
            if hrule == 2:
                heading = False
                hrule = 0
            continue
        elif line.startswith("•"):
            if heading:
                h = document.add_heading(line[1:], 1)
                heading = False
            else:
                sh = document.add_paragraph(line[1:])
                sh.style.paragraph_format.space_before = Pt(12)
                sh.style.paragraph_format.space_after = 0
            continue
        else:
            tabstops, words = parse(line)
            p = document.add_paragraph("\t" + "\t".join(words))
            p.style = "No Spacing"
            for tab in tabstops:
                _t = p.paragraph_format.tab_stops
                _t.add_tab_stop((Mm(tab * 1.8)), alignment=WD_TAB_ALIGNMENT.LEFT)
    document.save(join(output_directory, "sj_mepla.rep.docx"))


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("PATH", type=str, help="Path to mepla calculations folder")
    parser.add_argument(
        "-f",
        "--file",
        default="",
        choices=["", "dat", "rep"],
        type=str,
        help="Choose the desired file to parse. If none, both .dat and .rep are parsed",
    )
    args = parser.parse_args()
    PATH: str = args.PATH
    if not isdir(PATH):
        raise NotADirectoryError("The given path is not a valid directory")
    FILES: list[str] = [f for f in listdir(PATH) if isfile(join(PATH, f))]
    if args.file == "dat":
        if "sj_mepla.dat" in FILES:
            dat_file = "sj_mepla.dat"
            parse_dat(join(PATH, dat_file), PATH)
        else:
            raise FileNotFoundError("sj_mepla.dat file not found")
    elif args.file == "rep":
        if "sj_mepla.rep" in FILES:
            rep_file = "sj_mepla.rep"
            parse_rep(join(PATH, rep_file), PATH)
        else:
            raise FileNotFoundError("sj_mepla.rep file not found")
    else:
        all_files = [f for f in FILES if f.endswith(".rep") or f.endswith(".dat")]
        for _file in all_files:
            if _file == "sj_mepla.dat":
                parse_dat(join(PATH, _file), PATH)
            else:
                parse_rep(join(PATH, _file), PATH)
